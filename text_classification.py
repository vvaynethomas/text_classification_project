import nltk
import numpy as np
import pandas as pd
import os
import PyPDF2
import docx
import logging
import sys
import pickle
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import RandomizedSearchCV
from sklearn.model_selection import GridSearchCV
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score
from sklearn.model_selection import ShuffleSplit
from sklearn.decomposition import PCA
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.manifold import TSNE
from sklearn.naive_bayes import MultinomialNB
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.feature_selection import chi2
from sklearn import svm
from sklearn.neighbors import KNeighborsClassifier
from sklearn.linear_model import LogisticRegression
import random
import datetime

english_vocab = set(w.lower() for w in nltk.corpus.words.words())
nltk.download('words')
pd.set_option('mode.chained_assignment', None)

logging.basicConfig(filename='text_classifier2.log', level=logging.DEBUG)
logger = logging.getLogger(__name__)
fh = logging.FileHandler('text_classifier2.log', 'w+')
fh.setLevel(logging.DEBUG)
# create console handler with a higher log level
ch = logging.StreamHandler()
ch.setLevel(logging.INFO)
# create formatter and add it to the handlers
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
fh.setFormatter(formatter)
ch.setFormatter(formatter)
# add the handlers to the logger
logger.addHandler(fh)
logger.addHandler(ch)

punctuation_signs = list("?:!.,;")
nltk.download('stopwords')
stop_words = list(nltk.corpus.stopwords.words('english'))


def lemmatize_text(row):
    wordnet_lemmatizer = nltk.WordNetLemmatizer()
    text = row['Text']

    # Create an empty list containing lemmatized words
    lemmatized_list = []

    # Save the text and its words into an object
    text_words = text.split(" ")

    # Iterate through every word to lemmatize
    for word in text_words:
        lemmatized_list.append(wordnet_lemmatizer.lemmatize(word, pos="v"))
        
    # Join the list
    row['Text'] = " ".join(lemmatized_list)
    return row['Text']


def remove_stopwords(row):
    for stop_word in stop_words:
        regex_stopword = r"\b" + stop_word + r"\b"
        row['Text'] = row['Text'].replace(regex_stopword, '')
    return row['Text']


def extract_string_from_docx(doc_path):
    doc = None
    try:
        doc = docx.Document(doc_path)
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            return ' '.join(all_runs)

    except:
        return ''


def extract_string_from_pdf(pdf_path):
    # filename = os.path.basename(filepath)
    with open(pdf_path, 'rb') as pdf_file_obj:  # open current pdf
        try:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        except:
            print('unable to parse pdf')
            return ''
        if pdf_reader and not pdf_reader.isEncrypted:
            # print("I'm looking for a line in " + filename)
            try:
                number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
                if number_of_pages > 1:
                    text_in_lines = []
                    text_in_pages = []
                    for i in range(0, number_of_pages - 1):
                        page_obj = pdf_reader.getPage(i)
                        page_string = page_obj.extractText()
                        if 'content downloaded from' in page_string:
                            continue
                        else:
                            text_in_pages.append(page_string)
                    return ' '.join(text_in_pages)
            except:
                print("unparseable pdf")
                return ''


def extract_string_from_txt(txt_path):
    with open(txt_path, mode='r') as file:
        str_list = file.readlines()
        return ' '.join(str_list)


def get_document_paths(directory):
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    paths_for_potentials = []
    for path in file_paths:
        if path.endswith('.docx') or path.endswith('.pdf') or path.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials


def parse_text_column(df):
    df['Text'] = df.apply(parse_words, axis=1)
    return df


def parse_words(row):
    row['Text'] = clean_text(row)
    row['Text'] = lemmatize_text(row)
    row['Text'] = remove_stopwords(row)


def clean_text(row):
    row['Text'] = row['Text'].replace("\r", " ")
    row['Text'] = row['Text'].replace("\n", " ")
    row['Text'] = row['Text'].replace("    ", " ")
    row['Text'] = row['Text'].lower()
    for punct_sign in punctuation_signs:
        row['Text'] = row['Text'].replace(punct_sign, '')
    row['Text'] = row['Text'].replace("'s", "")
    return row['Text']


def df_to_corpus(df, text_column_name, label_column_name, title_column_name=None):
    logger.info("Begun raising corpus from dataframe")
    new_corpus = Corpus()
    for index, row in df.iterrows():
        new_doc = Document(text=row[text_column_name], label=row[label_column_name], name=row[title_column_name])
        new_corpus.add_document(new_doc)
    logger.info("Finished extracting corpus")
    return new_corpus


def doc_to_df(document):
    new_df = pd.DataFrame(columns=['Text', 'Label'])
    new_df.loc[0] = [document.text, document.label]
    return new_df


def csv_to_corpus(csv_path, text_column_name, label_column_name, title_column_name=None):
    df = pd.read_csv(csv_path)
    corpus = df_to_corpus(df, text_column_name, label_column_name, title_column_name)
    return corpus


def corpify(object, text_column_name=None, label_column_name=None, title_column_name=None):
    corpus = None
    if isinstance(object, Document):
        df = doc_to_df(Document)
        corpus = df_to_corpus(df, 'Text', 'Label')
    elif isinstance(object, pd.DataFrame):
        corpus = df_to_corpus(object, text_column_name, label_column_name, title_column_name)
    elif isinstance(object, str):
        if object.endswith('.csv'):
            corpus = csv_to_corpus(object, text_column_name, label_column_name, title_column_name)
        if object.endswith('.txt'):
            new_doc = Document(path=object)
            df = doc_to_df(new_doc)
            corpus = df_to_corpus(df, 'Text', 'Label')
    else:
        print("unsupported object type to be made into corpus")
        raise TypeError
    return corpus


class Document(object):
    '''Describes the Document object which has a text and a label in a dataframe form'''

    def __init__(self, path=None, name=None, label=None, text=None):
        logger.info("Beginning construction of Document object")
        logger.debug(f"Parameters passed: {path, name, label}")
        if path:
            self.path = path
            self.label = os.path.basename(os.path.dirname(self.path))
            self.name = os.path.basename(self.path).split('.')[0]
            self._extract_text()
        else:
            self.label = label.strip()
            self.text = text
            self.name = name
            self.path = str()
        # self.words = self._get_word_list()
        logger.info("Finished construction of Document object")

    def __repr__(self):
        return self.label + ' : ' + self.text

    def __str__(self):
        return f'Document object named {self.name} and labeled as {self.label}'

    def __add__(self, other):
        if type(other) == type(self):
            new_corpus = Corpus()
            new_corpus.documents.append(self)

    def _extract_text(self):
        if self.path.endswith('.docx'):
            self.text = extract_string_from_docx(self.path)
        elif self.path.endswith('.pdf'):
            self.text = extract_string_from_pdf(self.path)
        elif self.path.endswith('.txt'):
            self.text = extract_string_from_txt(self.path)

    def _get_word_list(self):
        text_in_words = nltk.tokenize.word_tokenize(self.text)
        final_words = []
        for word in text_in_words:
            if not word[0].isalnum() or not word[-1].isalnum():
                while word and not word[0].isalnum():
                    word = word[1:]
                while word and not word[-1].isalnum():
                    word = word[:-1]
            if len(word) <= 3 and word not in ["I'm", "I'd", "i'm", "i'd", "I’d", "I’m", 'i', 'I']:
                if word not in english_vocab and word.lower() not in english_vocab and word.upper() not in english_vocab:
                    word = ''
            if len(word) == 1 and word not in ['a', 'i', 'I']:
                word = ''
            if word:
                final_words.append(word)
        return final_words

    def _pos_list(self):
        return [tag for word, tag in nltk.pos_tag(self.words)]

    def word_count(self):
        return len(self.words)

    def noun_count(self):
        return len([word for word, tag in nltk.pos_tag(self.words) if 'NN' in tag])

    def get_unique(self):
        return sorted(set([word.lower() for word in self.words]))

    def count_unique(self):
        return len(self.get_unique())


class Corpus(object):
    def __init__(self, directory=None, name=None, pickle_bool=True):
        logger.info("Initializing Corpus object")
        logger.debug(f"Parameter passed: {str([type(i) for i in [directory, name, pickle_bool]])}")
        self.index = 0
        self.documents = []
        self.df = pd.DataFrame(columns=['Name', 'Text', 'Label'])
        if name:
            self.name = name
        if directory:
            doc_paths = get_document_paths(directory)
            for path in doc_paths:
                new_doc = Document(path)
                self.add_document(new_doc)
            if not name:
                self.name = os.path.basename(os.path.dirname(directory))

    def __add__(self, other):
        if type(self) == type(other):
            new_corpus = self
            for document in other.documents:
                new_corpus.add_document(document)
            return new_corpus
        elif type(other) == Document:
            self.add_document(other)
            return self

    def __iter__(self):
        return self

    def __next__(self):
        if self.index < len(self.documents):
            self.index += 1
            return self.documents[self.index - 1]
        else:
            raise StopIteration

    def add_document(self, document):
        logger.info("Adding document to Corpus")
        self.documents.append(document)
        # a_row = pd.Series([document.name, document.text, document.label])
        # row_df = pd.DataFrame([document.name, document.text, document.label])
        self.df = self.df.append({'Name': document.name, 'Text': document.text, 'Label': document.label},
                                 ignore_index=True)
        logger.info("Finished adding document to corpus")


class TextClassifier(object):
    nltk.download('punkt')
    nltk.download('wordnet')

    def __init__(self, inp,
                 delimiter=None,
                 method='tfidf',
                 model=None,
                 parameters=None,
                 plot_bool=False,
                 project=None,
                 pickle_bool=True,
                 threshold=3):
        logger.info("Initializing Classifier Object")
        logger.info("Checking for pickled object")
        pickle_out_name = str()

        if project and os.path.exists(project + '.p'):
            pickled_classifier = pickle.load(pickle_out_name)
            self.plot_bool = pickled_classifier.plot_bool
            self.models = pickled_classifier.models
            self.model_comparison = pickled_classifier.model_comparison
            self.all_data = pickled_classifier.all_data
            self.corpora = pickled_classifier.corpora
            self.method = pickled_classifier.method
            self.categories = pickled_classifier.categories
            self.tfidf = pickled_classifier.tfidf
            self.X_train = pickled_classifier.X_train
            self.X_test = pickled_classifier.X_test
            self.y_train = pickled_classifier.y_train
            self.y_test = pickled_classifier.y_test
            self.category_codes = pickled_classifier.category_codes

        else:
            self.plot_bool = plot_bool
            self.models = {}
            self.model_comparison = pd.DataFrame(columns=['Model', 'Training Set Accuracy', 'Test Set Accuracy'])
            # if parameters is None:
            #     parameters = [(None, None), None, None, None]
            self.all_data = pd.DataFrame(columns=['Name', 'Text', 'Label'])
            self.corpora = []
            if isinstance(inp, str):
                if os.path.exists(inp):
                    logger.debug(f'Path: {inp}, Method: {method}, Given Parameters: {parameters}')
                    if inp.endswith('.csv'):
                        logger.info("Processing file path as input")
                        if not project:
                            self.project = ''.join(list(os.path.basename(inp))[:-5])
                        # try:
                        if delimiter:
                            df = pd.read_csv(inp, delimiter=delimiter)
                        else:
                            df = pd.read_csv(inp)
                        print(df.columns)
                        if len(df.columns) > 2:
                            lcn = str(input("Enter label column name (case-sensitive) "))
                            tcn = str(input("Enter text column name (case-sensitive) "))
                        else:
                            col_1_len = df.iloc[:, 0].str.len()
                            col_2_len = df.iloc[:, 1].str.len()
                            if col_1_len.max > col_2_len.max:
                                tcn = df.columns[0]
                                lcn = df.columns[1]
                            else:
                                tcn = df.columns[1]
                                lcn = df.columns[0]
                            print(f'text column name: {tcn}')
                            print(f'label column name: {lcn}')
                        self.corpora.append(df_to_corpus(df, tcn, lcn))
                        self.all_data.append(df, ignore_index=True)
                        # except:
                        #     print("Problem converting file to Corpus, check header or delimiiter")
                        #     logger.debug("Failed to read and convert {inp}")
                    elif inp.endswith('.p'):
                        self.best_mod = pickle.load(open(inp, "rb"))
                    else:
                        logger.info("Processing Directory path as input")
                        if not project:
                            self.project = os.path.basename(os.path.dirname(inp))
                        else:
                            self.project = project
                        subdirs = [x[0] for x in os.walk(inp)]
                        for dir in subdirs:
                            cur_corp = Corpus(dir)
                            self.corpora.append(cur_corp)
                            self.all_data.append(cur_corp.df, ignore_index=True)
            elif isinstance(inp, Corpus):
                logger.info("Processing Corpus as input")
                if not project:
                    self.project = inp.name
                else:
                    self.project = project
                self.corpora.append(inp)
                self.all_data = inp.df
                logger.info("Finished absorbing corpus")
                logger.debug(f"ingested corpus has {len(self.all_data)} texts")
            # elif type(inp) == pd.DataFrame:
            #     logger.info("Procesing pandas datarame as input")
            #     this_corpus = df_to_corpus(inp,
            if not (isinstance(inp, str) and inp.endswith('.p')):
                # remove very rare (as determined by the threshold variable) labeled items
                counts = self.all_data['Label'].value_counts()
                self.all_data = self.all_data[self.all_data['Label'].isin(counts[counts > threshold].index)]
                logger.debug(f'Summary of pre-parsed text data to be classified: {self.all_data.describe()}')
                self.method = method
                logger.debug(f'method selected: {method}')
                self.categories = set(self.all_data['Label'].to_list())
                logger.debug(f'categories in corpus: {self.categories}')
                cat_range = [*range(0, len(self.categories))]
                cat_range = [str(i) for i in cat_range]
                self.category_codes = dict(zip(self.categories, cat_range))
                logger.debug(f'category codes: {self.category_codes}')
                self.all_data['Category_Code'] = self.all_data['Label']
                self.all_data = self.all_data.replace({'Category_Code': self.category_codes})
                self.X_train, self.X_test, self.y_train, self.y_test = train_test_split(self.all_data['Text'],
                                                                                        self.all_data['Category_Code'],
                                                                                        test_size=0.15,
                                                                                        random_state=8)
                logger.debug(f"y_train : {self.y_train}, y_test : {self.y_test}")
                if method == 'tfidf':
                    logger.info("Starting TF-IDF")
                    self._parse_texts()
                    if parameters:
                        ngram_range = parameters[0]
                        min_df = parameters[1]
                        max_df = parameters[2]
                        max_features = parameters[3]
                    else:
                        ngram_range = (1, 2)
                        min_df = threshold
                        max_df = 1.
                        max_features = 300

                    self.tfidf = TfidfVectorizer(encoding='utf-8',
                                                 ngram_range=ngram_range,
                                                 stop_words=None,
                                                 lowercase=False,
                                                 max_df=max_df,
                                                 min_df=min_df,
                                                 max_features=max_features,
                                                 norm='l2',
                                                 sublinear_tf=True)

                    self.features_train = self.tfidf.fit_transform(self.X_train).toarray()
                    self.labels_train = self.y_train

                    self.features_test = self.tfidf.transform(self.X_test).toarray()
                    self.labels_test = self.y_test
                    # # X_train
                    # with open('Pickles/X_train.pickle', 'wb') as output:
                    #     pickle.dump(X_train, output)
                    #
                    # # X_test
                    # with open('Pickles/X_test.pickle', 'wb') as output:
                    #     pickle.dump(X_test, output)
                    #
                    # # y_train
                    # with open('Pickles/y_train.pickle', 'wb') as output:
                    #     pickle.dump(y_train, output)
                    #
                    # # y_test
                    # with open('Pickles/y_test.pickle', 'wb') as output:
                    #     pickle.dump(y_test, output)
                    #
                    # # features_train
                    # with open('Pickles/features_train.pickle', 'wb') as output:
                    #     pickle.dump(features_train, output)
                    #
                    # # labels_train
                    # with open('Pickles/labels_train.pickle', 'wb') as output:
                    #     pickle.dump(labels_train, output)
                    #
                    # # features_test
                    # with open('Pickles/features_test.pickle', 'wb') as output:
                    #     pickle.dump(features_test, output)
                    #
                    # # labels_test
                    # with open('Pickles/labels_test.pickle', 'wb') as output:
                    #     pickle.dump(labels_test, output)
                    #
                    # # TF-IDF object
                    # with open('Pickles/tfidf.pickle', 'wb') as output:
                    #     pickle.dump(tfidf, output)
                    logger.info("tfidf complete, training and test features and labels stored")

            if model == 'random_forest':
                logger.info("beginning random forest classification")
                model, comp = self.random_forest()
                self.models['random_forest'] = model
                self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
                logger.info("random forest model saved")
                logger.debug(f'Model: {str(model)}')
                logger.debug(f'Accuracy: {str(comp)}')
                logger.debug(f'model comparison so far: {self.model_comparison}')

            if model == 'SVM':
                logger.info("beginning support vector machine classification")
                model, comp = self.SVM()
                self.models['SVM'] = model
                self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
                logger.info("supper vector machine classifier stored")
                logger.debug(f'Model: {str(model)}')
                logger.debug(f'Accuracy: {str(comp)}')
                logger.debug(f'model comparison so far: {self.model_comparison}')

            if model == 'KNN':
                logger.info("beginning k-nearest neighbors classification")
                model, comp = self.KNN()
                self.models['KNN'] = model
                self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
                logger.info("stored k-nearest neighbors classifier")
                logger.debug(f'Model: {str(model)}')
                logger.debug(f'Accuracy: {str(comp)}')
                logger.debug(f'model comparison so far: {self.model_comparison}')

            if model == 'MNB':
                model, comp = self.MNB()
                self.models['MNB'] = model
                self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
                logger.debug(f'Model: {str(model)}')
                logger.debug(f'Accuracy: {str(comp)}')
                logger.debug(f'model comparison so far: {self.model_comparison}')

            if model == 'MLR':
                model, comp = self.MLR()
                self.models['MLR'] = model
                self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
                logger.debug(f'Model: {str(model)}')
                logger.debug(f'Accuracy: {str(comp)}')
                logger.debug(f'model comparison so far: {self.model_comparison}')

            if model == 'GBM':
                model, comp = self.GBM()
                self.models['GBM'] = model
                self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
                logger.debug(f'Model: {str(model)}')
                logger.debug(f'Accuracy: {str(comp)}')
                logger.debug(f'model comparison so far: {self.model_comparison}')

            if model == 'best':
                logger.info('running all models to find the best one')
                model, comp = self.best()

            if model:
                df_summary = self.model_comparison
                print(df_summary.sort_values('Test Set Accuracy', ascending=False))

                if model == 'best':
                    not_overfitted = df_summary[df_summary['Training Set Accuracy'] <= .98]
                    best_Test = not_overfitted[
                        not_overfitted['Test Set Accuracy'] == not_overfitted['Test Set Accuracy'].max()]
                    self.best_mod_key = str(best_Test['Model'])
                else:
                    self.best_mod_key = str(df_summary['Model'])
                self.best_mod = self.models[self.best_mod_key]

            if not os.path.exists(pickle_out_name) or pickle_bool:
                pickle_out_name = self.project + '.p'
                pickle.dump(self, open(pickle_out_name, 'wb'))

    def _parse_texts(self):
        self.all_data = parse_text_column(self.all_data)

    # show what words and bigrams are most correlated to which labels
    def cor_unigrams_and_bigrams(self):
        for label, category_id in sorted(self.category_codes.items()):
            features_chi2 = chi2(self.features_train, self.labels_train == category_id)
            indices = np.argsort(features_chi2[0])
            feature_names = np.array(self.tfidf.get_feature_names())[indices]
            unigrams = [v for v in feature_names if len(v.split(' ')) == 1]
            bigrams = [v for v in feature_names if len(v.split(' ')) == 2]
            print("# '{}' category:".format(label))
            print("  . Most correlated unigrams:\n. {}".format('\n. '.join(unigrams[-5:])))
            print("  . Most correlated bigrams:\n. {}".format('\n. '.join(bigrams[-2:])))
            print("")

    # produce random forest model
    def random_forest(self):
        logger.info("Beginning Random Forest Modeling")
        # # initialize random forest classifier
        # rf_0 = RandomForestClassifier(random_state=8)
        # logger.debug(f"Current parameters: \n {rf_0.get_params()}")

        # n_estimators
        n_estimators = [int(x) for x in np.linspace(start=200, stop=1000, num=5)]

        # max_features
        max_features = ['auto', 'sqrt']

        # max_depth
        max_depth = [int(x) for x in np.linspace(20, 100, num=5)]
        max_depth.append(None)

        # min_samples_split
        min_samples_split = [2, 5, 10]

        # min_samples_leaf
        min_samples_leaf = [1, 2, 4]

        # bootstrap
        bootstrap = [True, False]

        # Create the random grid
        random_grid = {'n_estimators': n_estimators,
                       'max_features': max_features,
                       'max_depth': max_depth,
                       'min_samples_split': min_samples_split,
                       'min_samples_leaf': min_samples_leaf,
                       'bootstrap': bootstrap}
        logger.debug(f"random grid \n {random_grid}")

        # contruct base model
        rfc = RandomForestClassifier(random_state=8)

        # Definition of the random search
        random_search = RandomizedSearchCV(estimator=rfc,
                                           param_distributions=random_grid,
                                           n_iter=50,
                                           scoring='accuracy',
                                           cv=3,
                                           verbose=1,
                                           random_state=8)

        # Fit the random search model
        random_search.fit(self.features_train, self.labels_train)

        print("The best hyperparameters from Random Search are:")
        print(random_search.best_params_)
        logger.debug(f"random forest random search best hyperparameters: {random_search.best_params_}")
        print("")
        print("The mean accuracy of a model with these hyperparameters is:")
        print(random_search.best_score_)

        # Create the parameter grid based on the results of random search
        bootstrap = [random_search.best_params_['bootstrap']]
        max_depth = [random_search.best_params_['max_depth']]
        max_features = [random_search.best_params_['max_features']]
        min_samples_leaf = [random_search.best_params_['min_samples_leaf']]
        min_samples_split = [random_search.best_params_['min_samples_split']]
        n_estimators = [random_search.best_params_['n_estimators']]

        param_grid = {
            'bootstrap': bootstrap,
            'max_depth': max_depth,
            'max_features': max_features,
            'min_samples_leaf': min_samples_leaf,
            'min_samples_split': min_samples_split,
            'n_estimators': n_estimators
        }

        # Manually create the splits in CV in order to be able to fix a random_state (GridSearchCV doesn't have that argument)
        cv_sets = ShuffleSplit(n_splits=3, test_size=.33, random_state=8)

        # Instantiate the grid search model
        grid_search = GridSearchCV(estimator=rfc,
                                   param_grid=param_grid,
                                   scoring='accuracy',
                                   cv=cv_sets,
                                   verbose=1)

        # Fit the grid search to the data
        grid_search.fit(self.features_train, self.labels_train)

        print("The best hyperparameters from Grid Search are:")
        print(grid_search.best_params_)
        print("")
        print("The mean accuracy of a model with these hyperparameters is:")
        print(grid_search.best_score_)

        best_rfc = grid_search.best_estimator_
        print(best_rfc)
        best_rfc.fit(self.features_train, self.labels_train)

        rfc_pred = best_rfc.predict(self.features_test)

        # Training accuracy
        print("The training accuracy is: ")
        print(accuracy_score(self.labels_train, best_rfc.predict(self.features_train)))

        # Test accuracy
        print("The test accuracy is: ")
        print(accuracy_score(self.labels_test, rfc_pred))

        # Classification report
        print("Classification report")
        print(classification_report(self.labels_test, rfc_pred))

        if self.plot_bool:
            aux_df = self.all_data[['Label', 'Category_Code']].drop_duplicates().sort_values('Category_Code')
            conf_matrix = confusion_matrix(self.labels_test, rfc_pred)
            plt.figure(figsize=(12.8, 6))
            sns.heatmap(conf_matrix,
                        annot=True,
                        xticklabels=aux_df['Label'].values,
                        yticklabels=aux_df['Label'].values,
                        cmap="Blues")
            plt.ylabel('Predicted')
            plt.xlabel('Actual')
            plt.title('Confusion matrix')
            plt.savefig(f'rf_heatmap_{self.project}.png')
            plt.show()

        d = {
            'Model': 'Random_forest',
            'Training Set Accuracy': accuracy_score(self.labels_train, best_rfc.predict(self.features_train)),
            'Test Set Accuracy': accuracy_score(self.labels_test, rfc_pred)
        }

        df_models_rfc = pd.DataFrame(d, index=[0])
        print(df_models_rfc)
        return best_rfc, df_models_rfc

    # produce support vector machine model
    def SVM(self):
        logger.info("Initializing Support Vector Machine modeling")
        # # hyperparameter tuning with randomized search cross validation
        #
        # # C
        # C = [.0001, .001, .01]
        #
        # # gamma
        # gamma = [.0001, .001, .01, .1, 1, 10, 100]
        #
        # # degree
        # degree = [1, 2, 3, 4, 5]
        #
        # # kernel
        # kernel = ['linear', 'rbf', 'poly']
        #
        # # probability
        # probability = [True]
        #
        # # Create the random grid
        # random_grid = {'C': C,
        #                'kernel': kernel,
        #                'gamma': gamma,
        #                'degree': degree,
        #                'probability': probability
        #                }
        #
        # svc = svm.SVC(random_state=8)
        #
        # # Definition of the random search
        # random_search = RandomizedSearchCV(estimator=svc,
        #                                    param_distributions=random_grid,
        #                                    n_iter=50,
        #                                    scoring='accuracy',
        #                                    cv=3,
        #                                    verbose=1,
        #                                    random_state=8)
        #
        # # Fit the random search model
        # random_search.fit(self.features_train, self.labels_train)
        #
        # print("The best hyperparameters from Random Search are:")
        # print(random_search.best_params_)
        # print("")
        # print("The mean accuracy of a model with these hyperparameters is:")
        # print(random_search.best_score_)

        # svm hyperparameter tuning with grid search cross validation

        # Create the parameter grid based on the results of random search
        C = [.0001, .001, .01, .1]
        degree = [3, 4, 5]
        gamma = [1, 10, 100]
        probability = [True]

        param_grid = [
            {'C': C, 'kernel': ['linear'], 'probability': probability},
            {'C': C, 'kernel': ['poly'], 'degree': degree, 'probability': probability},
            {'C': C, 'kernel': ['rbf'], 'gamma': gamma, 'probability': probability}
        ]

        # Create a base model
        svc = svm.SVC(random_state=8)

        # Manually create the splits in CV in order to be able to fix a random_state (GridSearchCV doesn't have that argument)
        cv_sets = ShuffleSplit(n_splits=3, test_size=.33, random_state=8)

        # Instantiate the grid search model
        grid_search = GridSearchCV(estimator=svc,
                                   param_grid=param_grid,
                                   scoring='accuracy',
                                   cv=cv_sets,
                                   verbose=1)

        # Fit the grid search to the data
        grid_search.fit(self.features_train, self.labels_train)

        print("The best hyperparameters from Grid Search are:")
        print(grid_search.best_params_)
        print("")
        print("The mean accuracy of a model with these hyperparameters is:")
        print(grid_search.best_score_)

        best_svc = grid_search.best_estimator_
        best_svc.fit(self.features_train, self.labels_train)
        svc_pred = best_svc.predict(self.features_test)

        # Training accuracy
        print("The training accuracy is: ")
        print(accuracy_score(self.labels_train, best_svc.predict(self.features_train)))

        # Test accuracy
        print("The test accuracy is: ")
        print(accuracy_score(self.labels_test, svc_pred))

        # Classification report
        print("Classification report")
        print(classification_report(self.labels_test, svc_pred))

        if self.plot_bool:
            aux_df = self.all_data[['Label', 'Category_Code']].drop_duplicates().sort_values('Category_Code')
            conf_matrix = confusion_matrix(self.labels_test, svc_pred)
            plt.figure(figsize=(12.8, 6))
            sns.heatmap(conf_matrix,
                        annot=True,
                        xticklabels=aux_df['Label'].values,
                        yticklabels=aux_df['Label'].values,
                        cmap="Blues")
            plt.ylabel('Predicted')
            plt.xlabel('Actual')
            plt.title('Confusion matrix')
            plt.savefig(f'svm_heatmap_{self.project}')
            plt.show()

        d = {
            'Model': 'SVM',
            'Training Set Accuracy': accuracy_score(self.labels_train, best_svc.predict(self.features_train)),
            'Test Set Accuracy': accuracy_score(self.labels_test, svc_pred)
        }

        df_models_svc = pd.DataFrame(d, index=[0])
        return best_svc, df_models_svc

    # produce KNN model
    def KNN(self):
        n_neighbors = [int(x) for x in np.linspace(start=1, stop=200, num=100)]

        param_grid = {'n_neighbors': n_neighbors}

        # Create a base model
        knnc = KNeighborsClassifier()

        # Manually create the splits in CV in order to be able to fix a random_state (GridSearchCV doesn't have that argument)
        cv_sets = ShuffleSplit(n_splits=3, test_size=.15, random_state=8)

        # Instantiate the grid search model
        grid_search = GridSearchCV(estimator=knnc,
                                   param_grid=param_grid,
                                   scoring='accuracy',
                                   cv=cv_sets,
                                   verbose=1)

        # Fit the grid search to the data
        grid_search.fit(self.features_train, self.labels_train)

        print("The best hyperparameters from Grid Search are:")
        print(grid_search.best_params_)
        print("")
        print("The mean accuracy of a model with these hyperparameters is:")
        print(grid_search.best_score_)

        best_knnc = grid_search.best_estimator_

        best_knnc.fit(self.features_train, self.labels_train)
        knnc_pred = best_knnc.predict(self.features_test)

        # Training accuracy
        print("The training accuracy is: ")
        print(accuracy_score(self.labels_train, best_knnc.predict(self.features_train)))

        # Test accuracy
        print("The test accuracy is: ")
        print(accuracy_score(self.labels_test, knnc_pred))

        # Classification report
        print("Classification report")
        print(classification_report(self.labels_test, knnc_pred))
        if self.plot_bool:
            aux_df = self.all_data[['Label', 'Category_Code']].drop_duplicates().sort_values('Category_Code')
            conf_matrix = confusion_matrix(self.labels_test, knnc_pred)
            plt.figure(figsize=(12.8, 6))
            sns.heatmap(conf_matrix,
                        annot=True,
                        xticklabels=aux_df['Label'].values,
                        yticklabels=aux_df['Label'].values,
                        cmap="Blues")
            plt.ylabel('Predicted')
            plt.xlabel('Actual')
            plt.title('Confusion matrix')
            plt.savefig(f'knn_heatmap_{self.project}.png')
            plt.show()

        d = {
            'Model': 'KNN',
            'Training Set Accuracy': accuracy_score(self.labels_train, best_knnc.predict(self.features_train)),
            'Test Set Accuracy': accuracy_score(self.labels_test, knnc_pred)
        }

        df_models_knnc = pd.DataFrame(d, index=[0])
        return best_knnc, df_models_knnc

    # produce multinomial naive bayes model
    def MNB(self):
        mnbc = MultinomialNB()
        mnbc.fit(X=self.features_train, y=self.labels_train)
        mnbc_pred = mnbc.predict(self.features_test)

        # Training accuracy
        print("The training accuracy is: ")
        print(accuracy_score(self.labels_train, mnbc.predict(self.features_train)))

        # Test accuracy
        print("The test accuracy is: ")
        print(accuracy_score(self.labels_test, mnbc_pred))

        # Classification report
        print("Classification report")
        print(classification_report(self.labels_test, mnbc_pred))

        if self.plot_bool:
            aux_df = self.all_data[['Label', 'Category_Code']].drop_duplicates().sort_values('Category_Code')
            conf_matrix = confusion_matrix(self.labels_test, mnbc_pred)
            plt.figure(figsize=(12.8, 6))
            sns.heatmap(conf_matrix,
                        annot=True,
                        xticklabels=aux_df['Label'].values,
                        yticklabels=aux_df['Label'].values,
                        cmap="Blues")
            plt.ylabel('Predicted')
            plt.xlabel('Actual')
            plt.title('Confusion matrix')
            plt.savefig(f'mnb_heatmap_{self.project}.png')
            plt.show()

        d = {
            'Model': 'MNB',
            'Training Set Accuracy': accuracy_score(self.labels_train, mnbc.predict(self.features_train)),
            'Test Set Accuracy': accuracy_score(self.labels_test, mnbc_pred)
        }

        df_models_mnbc = pd.DataFrame(d, index=[0])
        return mnbc, df_models_mnbc

    # produce multinomial logistic regression model
    def MLR(self):
        C = [float(x) for x in np.linspace(start=0.6, stop=1, num=10)]
        multi_class = ['multinomial']
        solver = ['sag']
        class_weight = ['balanced']
        penalty = ['l2']

        param_grid = {'C': C,
                      'multi_class': multi_class,
                      'solver': solver,
                      'class_weight': class_weight,
                      'penalty': penalty}

        # Create a base model
        lrc = LogisticRegression(random_state=8)

        # Manually create the splits in CV in order to be able to fix a random_state (GridSearchCV doesn't have that argument)
        cv_sets = ShuffleSplit(n_splits=3, test_size=.33, random_state=8)

        # Instantiate the grid search model
        grid_search = GridSearchCV(estimator=lrc,
                                   param_grid=param_grid,
                                   scoring='accuracy',
                                   cv=cv_sets,
                                   verbose=1)

        # Fit the grid search to the data
        grid_search.fit(self.features_train, self.labels_train)

        print("The best hyperparameters from Grid Search are:")
        print(grid_search.best_params_)
        print("")
        print("The mean accuracy of a model with these hyperparameters is:")
        print(grid_search.best_score_)

        best_lrc = grid_search.best_estimator_

        best_lrc.fit(self.features_train, self.labels_train)
        lrc_pred = best_lrc.predict(self.features_test)

        # Training accuracy
        print("The training accuracy is: ")
        print(accuracy_score(self.labels_train, best_lrc.predict(self.features_train)))

        # Test accuracy
        print("The test accuracy is: ")
        print(accuracy_score(self.labels_test, lrc_pred))

        # Classification report
        print("Classification report")
        print(classification_report(self.labels_test, lrc_pred))

        if self.plot_bool:
            aux_df = self.all_data[['Label', 'Category_Code']].drop_duplicates().sort_values('Category_Code')
            conf_matrix = confusion_matrix(self.labels_test, lrc_pred)
            plt.figure(figsize=(12.8, 6))
            sns.heatmap(conf_matrix,
                        annot=True,
                        xticklabels=aux_df['Label'].values,
                        yticklabels=aux_df['Label'].values,
                        cmap="Blues")
            plt.ylabel('Predicted')
            plt.xlabel('Actual')
            plt.title('Confusion matrix')
            plt.savefig(f'MLR_heatmap_{self.project}.png')
            plt.show()

        d = {
            'Model': 'MLR',
            'Training Set Accuracy': accuracy_score(self.labels_train, best_lrc.predict(self.features_train)),
            'Test Set Accuracy': accuracy_score(self.labels_test, lrc_pred)
        }

        df_models_lrc = pd.DataFrame(d, index=[0])
        return best_lrc, df_models_lrc

    # produce gradient boosting machine model
    def GBM(self):
        logger.info("Starting gradient boosted machine classification")
        from sklearn.ensemble import GradientBoostingClassifier
        # n_estimators
        n_estimators = [200, 800]

        # max_features
        max_features = ['auto', 'sqrt']

        # max_depth
        max_depth = [10, 40]
        max_depth.append(None)

        # min_samples_split
        min_samples_split = [10, 30, 50]

        # min_samples_leaf
        min_samples_leaf = [1, 2, 4]

        # learning rate
        learning_rate = [.1, .5]

        # subsample
        subsample = [.5, 1.]

        # Create the random grid
        random_grid = {'n_estimators': n_estimators,
                       'max_features': max_features,
                       'max_depth': max_depth,
                       'min_samples_split': min_samples_split,
                       'min_samples_leaf': min_samples_leaf,
                       'learning_rate': learning_rate,
                       'subsample': subsample}

        # First create the base model to tune
        gbc = GradientBoostingClassifier(random_state=8)

        # Definition of the random search
        random_search = RandomizedSearchCV(estimator=gbc,
                                           param_distributions=random_grid,
                                           n_iter=50,
                                           scoring='accuracy',
                                           cv=3,
                                           verbose=1,
                                           random_state=8)
        logger.info("Fitting random search model")
        # Fit the random search model
        random_search.fit(self.features_train, self.labels_train)
        best_dict = random_search.best_params_
        logger.debug(f"random search parameters = {best_dict}")
        # Create the parameter grid based on the results of random search
        max_depth = [best_dict['max_depth']]
        max_features = [best_dict['max_features']]
        min_samples_leaf = [best_dict['min_samples_leaf']]
        min_samples_split = [best_dict['min_samples_split']]
        n_estimators = [best_dict['n_estimators']]
        learning_rate = [best_dict['learning_rate']]
        subsample = [best_dict['subsample']]

        param_grid = {
            'max_depth': max_depth,
            'max_features': max_features,
            'min_samples_leaf': min_samples_leaf,
            'min_samples_split': min_samples_split,
            'n_estimators': n_estimators,
            'learning_rate': learning_rate,
            'subsample': subsample

        }

        # Create a base model
        gbc = GradientBoostingClassifier(random_state=8)

        # Manually create the splits in CV in order to be able to fix a random_state (GridSearchCV doesn't have that argument)
        cv_sets = ShuffleSplit(n_splits=3, test_size=.33, random_state=8)

        logger.info("running grid search with random search parameters")

        # Instantiate the grid search model
        grid_search = GridSearchCV(estimator=gbc,
                                   param_grid=param_grid,
                                   scoring='accuracy',
                                   cv=cv_sets,
                                   verbose=1)

        # Fit the grid search to the data
        grid_search.fit(self.features_train, self.labels_train)

        print("The best hyperparameters from Grid Search are:")
        print(grid_search.best_params_)
        print("")
        print("The mean accuracy of a model with these hyperparameters is:")
        print(grid_search.best_score_)

        best_gbc = grid_search.best_estimator_

        best_gbc.fit(self.features_train, self.labels_train)
        gbc_pred = best_gbc.predict(self.features_test)

        # Training accuracy
        print("The training accuracy is: ")
        print(accuracy_score(self.labels_train, best_gbc.predict(self.features_train)))

        # Test accuracy
        print("The test accuracy is: ")
        print(accuracy_score(self.labels_test, gbc_pred))

        # Classification report
        print("Classification report")
        print(classification_report(self.labels_test, gbc_pred))
        if self.plot_bool:
            aux_df = self.all_data[['Label', 'Category_Code']].drop_duplicates().sort_values('Category_Code')
            conf_matrix = confusion_matrix(self.labels_test, gbc_pred)
            plt.figure(figsize=(12.8, 6))
            sns.heatmap(conf_matrix,
                        annot=True,
                        xticklabels=aux_df['Label'].values,
                        yticklabels=aux_df['Label'].values,
                        cmap="Blues")
            plt.ylabel('Predicted')
            plt.xlabel('Actual')
            plt.title('Confusion matrix')
            plt.savefig(f'GBM_heatmap_{self.project}.png')
            plt.show()

        d = {
            'Model': 'GBM',
            'Training Set Accuracy': accuracy_score(self.labels_train, best_gbc.predict(self.features_train)),
            'Test Set Accuracy': accuracy_score(self.labels_test, gbc_pred)
        }

        df_models_gbc = pd.DataFrame(d, index=[0])
        return best_gbc, df_models_gbc

    # find best model
    def best(self):
        logger.info("beginning random forest classification")
        model, comp = self.random_forest()
        self.models['random_forest'] = model
        self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
        logger.info("random forest model saved")
        logger.debug(f'Model: {str(model)}')
        logger.debug(f'Accuracy: {str(comp)}')
        logger.debug(f'model comparison so far: {self.model_comparison}')

        logger.info("beginning support vector machine classification")
        model, comp = self.SVM()
        self.models['SVM'] = model
        self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
        logger.info("supper vector machine classifier stored")
        logger.debug(f'Model: {str(model)}')
        logger.debug(f'Accuracy: {str(comp)}')
        logger.debug(f'model comparison so far: {self.model_comparison}')

        logger.info("beginning k-nearest neighbors classification")
        model, comp = self.KNN()
        self.models['KNN'] = model
        self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
        logger.info("stored k-nearest neighbors classifier")
        logger.debug(f'Model: {str(model)}')
        logger.debug(f'Accuracy: {str(comp)}')
        logger.debug(f'model comparison so far: {self.model_comparison}')

        model, comp = self.MNB()
        self.models['MNB'] = model
        self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
        logger.debug(f'Model: {str(model)}')
        logger.debug(f'Accuracy: {str(comp)}')
        logger.debug(f'model comparison so far: {self.model_comparison}')

        model, comp = self.MLR()
        self.models['MLR'] = model
        self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
        logger.debug(f'Model: {str(model)}')
        logger.debug(f'Accuracy: {str(comp)}')
        logger.debug(f'model comparison so far: {self.model_comparison}')

        model, comp = self.GBM()
        self.models['GBM'] = model
        self.model_comparison = self.model_comparison.append(comp, ignore_index=True)
        logger.debug(f'Model: {str(model)}')
        logger.debug(f'Accuracy: {str(comp)}')
        logger.debug(f'model comparison so far: {self.model_comparison}')

        df_summary = self.model_comparison
        print(df_summary.sort_values('Test Set Accuracy', ascending=False))

        not_overfitted = df_summary[df_summary['Training Set Accuracy'] <= .98]
        best_Test = not_overfitted[
            not_overfitted['Test Set Accuracy'] == not_overfitted['Test Set Accuracy'].max()]
        self.best_mod_key = str(best_Test['Model'])
        self.best_mod = self.models[self.best_mod_key]
        return self.best_mod, df_summary

    # dimensionality reduction
    def plot_dim_red(self, model, features, labels, n_components=2):
        # Creation of the model
        if (model == 'PCA'):
            mod = PCA(n_components=n_components)
            title = "PCA decomposition"  # for the plot

        elif (model == 'TSNE'):
            mod = TSNE(n_components=2)
            title = "t-SNE decomposition"

        else:
            return "Error"

        # Fit and transform the features
        principal_components = mod.fit_transform(features)

        # Put them into a dataframe
        df_features = pd.DataFrame(data=principal_components,
                                   columns=['PC1', 'PC2'])

        # Now we have to paste each row's label and its meaning
        # Convert labels array to df
        df_labels = pd.DataFrame(data=labels,
                                 columns=['label'])

        df_full = pd.concat([df_features, df_labels], axis=1)
        df_full['label'] = df_full['label'].astype(str)

        # Get labels name
        category_names = self.categories

        # And map labels
        df_full['label_name'] = df_full['label']
        df_full = df_full.replace({'label_name': category_names})

        # Plot
        plt.figure(figsize=(10, 10))
        sns_plot = sns.scatterplot(x='PC1',
                                   y='PC2',
                                   hue="label_name",
                                   data=df_full,
                                   palette=["red", "pink", "royalblue", "greenyellow", "lightseagreen"],
                                   alpha=.7).set_title(title)
        sns_plot.savefig(f"dimension_reduction_{self.project}.png")

    def dim_red(self):
        features = np.concatenate((self.features_train, self.features_test), axis=0)
        labels = np.concatenate((self.labels_train, self.labels_test), axis=0)
        self.plot_dim_red("PCA", features=features, labels=labels, n_components=2)
        self.plot_dim_red("TSNE", features=features, labels=labels, n_components=2)

    def mod_test(self, ):
        predictions = self.best_mod.predict(self.features_test)
        # Indexes of the test set
        index_X_test = self.X_test.index

        # We get them from the original df
        df_test = self.all_data.loc[index_X_test]

        # Add the predictions
        df_test['Prediction'] = predictions

        # Clean columns
        df_test = df_test[['Text', 'Label', 'Category_Code', 'Prediction']]

        # Decode
        df_test['Category_Predicted'] = df_test['Prediction']
        df_test = df_test.replace({'Category_Predicted': self.categories})

        # Clean columns again
        df_test = df_test[['Content', 'Category', 'Category_Predicted']]

        condition = (df_test['Label'] != df_test['Category_Predicted'])

        df_misclassified = df_test[condition]
        return df_misclassified

    def mod_predict(self, text_to_classify):
        def create_features_from_text(text_object):
            df = pd.DataFrame(columns=['Text'])
            if type(text_object) == str:
                df.loc[0] = text_to_classify
            elif type(text_object) == Document:
                df.loc[0] = text_object.text

            df = parse_text_column(df)

            # TF-IDF
            features = self.tfidf.transform(df).toarray()
            return features

        def get_category_name(category_id):
            for category, id in self.category_codes.items():
                if id == category_id:
                    return category

        def predict_from_text(text):
            # Predict using the input model
            prediction_best = self.best_mod.predict(create_features_from_text(text))[0]
            prediction_best_proba = self.best_mod.predict_proba(create_features_from_text(text))[0]

            # Return result
            category_best = get_category_name(prediction_best)
            if not self.best_mod_key:
                self.best_mod_key = 'given'
            print(f"The predicted category using the {self.best_mod_key} model is {category_best}.")
            print("The conditional probability is: %a" % (prediction_best_proba.max() * 100))
            return category_best

        prediction = predict_from_text(text_to_classify)
        return prediction


def main():
    this_corpus = corpify(object='story.data.csv',
                          text_column_name='content',
                          label_column_name='Author',
                          title_column_name='Title')
    this_classifier = TextClassifier(inp=this_corpus, project='Short_stories')
    return this_classifier
