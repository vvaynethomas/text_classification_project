#! Python 3 -- text_extraction.py contains functions for extracting strings from different types of documents

import os
import PyPDF2
import docx
import pandas as pd
import nltk
import re
import logging

pd.set_option('mode.chained_assignment', None)

logging.basicConfig(filename='text_extraction.log', level=logging.DEBUG)
logger = logging.getLogger(__name__)
fh = logging.FileHandler('text_extraction.log', 'w+')
fh.setLevel(logging.DEBUG)
# create console handler with a higher log level
ch = logging.StreamHandler()
ch.setLevel(logging.INFO)
# create formatter and add it to the handlers
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
fh.setFormatter(formatter)
ch.setFormatter(formatter)
# add the handlers to the logger
logger.addHandler(fh)
logger.addHandler(ch)

english_vocab = set(w.lower() for w in nltk.corpus.words.words())

caps = "([A-Z])"
prefixes = "(Mr|St|Mrs|Ms|Dr)[.]"
suffixes = "(Inc|Ltd|Jr|Sr|Co)"
starters = "(Mr|Mrs|Ms|Dr|He\s|She\s|It\s|They\s|Their\s|Our\s|We\s|But\s|However\s|That\s|This\s|Wherever)"
acronyms = "([A-Z][.][A-Z][.](?:[A-Z][.])?)"
websites = "[.](com|net|org|io|gov)"


def extract_string_from_docx(doc_path):
    doc = None
    try:
        doc = docx.Document(doc_path)
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            return ' '.join(all_runs)

    except:
        return ''


def extract_string_from_pdf(pdf_path):
    # filename = os.path.basename(filepath)
    with open(pdf_path, 'rb') as pdf_file_obj:  # open current pdf
        try:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        except:
            print('unable to parse pdf')
            return ''
        if pdf_reader and not pdf_reader.isEncrypted:
            # print("I'm looking for a line in " + filename)
            try:
                number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
                if number_of_pages > 1:
                    text_in_lines = []
                    text_in_pages = []
                    for i in range(0, number_of_pages - 1):
                        page_obj = pdf_reader.getPage(i)
                        page_string = page_obj.extractText()
                        if 'content downloaded from' in page_string:
                            continue
                        else:
                            text_in_pages.append(page_string)
                    return ' '.join(text_in_pages)
            except:
                print("unparseable pdf")
                return ''


def extract_string_from_txt(txt_path):
    with open(txt_path, mode='r') as file:
        str_list = file.readlines()
        return ' '.join(str_list)


def get_document_paths(directory):
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    paths_for_potentials = []
    for path in file_paths:
        if path.endswith('.docx') or path.endswith('.pdf') or path.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials


def df_to_corpus(df, text_column_name, label_column_name, title_column_name=None):
    logger.info("Begun raising corpus from dataframe")
    new_corpus = Corpus()
    for index, row in df.iterrows():
        new_doc = Document(text=row[text_column_name], label=row[label_column_name], name=row[title_column_name])
        new_corpus.add_document(new_doc)
    logger.info("Finished extracting corpus")
    return new_corpus


def doc_to_df(document):
    new_df = pd.DataFrame(columns=['Text', 'Label'])
    new_df.loc[0] = [document.text, document.label]
    return new_df


def csv_to_corpus(csv_path, text_column_name, label_column_name, title_column_name=None):
    df = pd.read_csv(csv_path)
    corpus = df_to_corpus(df, text_column_name, label_column_name, title_column_name)
    return corpus


def corpify(object, text_column_name=None, label_column_name=None, title_column_name=None):
    if isinstance(object, Document):
        df = doc_to_df(Document)
        corpus = df_to_corpus(df, 'Text', 'Label')
    if isinstance(object, pd.DataFrame):
        corpus = df_to_corpus(df, text_column_name, label_column_name, title_column_name)
    if isinstance(object, str):
        if object.endswith('.csv'):
            corpus = csv_to_corpus(object, text_column_name, label_column_name, title_column_name)
        if object.endswith('.txt'):
            new_doc = Document(path=object)
            df = doc_to_df(new_doc)
            corpus = df_to_corpus(df, 'Text', 'Label')
    return corpus


class Document(object):
    '''Describes the Document object which has a text and a label in a dataframe form'''

    def __init__(self, path=None, name=None, label=None, text=None):
        logger.info("Beginning construction of Document object")
        logger.debug(f"Parameters passed: {path, name, label}")
        if path:
            self.path = path
            self.label = os.path.basename(os.path.dirname(self.path))
            self.name = os.path.basename(self.path).split('.')[0]
            self._extract_text()
        else:
            self.label = label.strip()
            self.text = text
            self.name = name
            self.path = str()
        # self.words = self._get_word_list()
        logger.info("Finished construction of Document object")

    def __repr__(self):
        return self.label + ' : ' + self.text

    def __str__(self):
        return f'Document object named {self.name} and labeled as {self.label}'

    def __add__(self, other):
        if type(other) == type(self):
            new_corpus = Corpus()
            new_corpus.documents.append(self)

    def _extract_text(self):
        if self.path.endswith('.docx'):
            self.text = extract_string_from_docx(self.path)
        elif self.path.endswith('.pdf'):
            self.text = extract_string_from_pdf(self.path)
        elif self.path.endswith('.txt'):
            self.text = extract_string_from_txt(self.path)

    def _get_word_list(self):
        text_in_words = nltk.tokenize.word_tokenize(self.text)
        final_words = []
        for word in text_in_words:
            if not word[0].isalnum() or not word[-1].isalnum():
                while word and not word[0].isalnum():
                    word = word[1:]
                while word and not word[-1].isalnum():
                    word = word[:-1]
            if len(word) <= 3 and word not in ["I'm", "I'd", "i'm", "i'd", "I’d", "I’m", 'i', 'I']:
                if word not in english_vocab and word.lower() not in english_vocab and word.upper() not in english_vocab:
                    word = ''
            if len(word) == 1 and word not in ['a', 'i', 'I']:
                word = ''
            if word:
                final_words.append(word)
        return final_words

    def _pos_list(self):
        return [tag for word, tag in nltk.pos_tag(self.words)]

    def word_count(self):
        return len(self.words)

    def noun_count(self):
        return len([word for word, tag in nltk.pos_tag(self.words) if 'NN' in tag])

    def get_unique(self):
        return sorted(set([word.lower() for word in self.words]))

    def count_unique(self):
        return len(self.get_unique())


class Corpus(object):
    def __init__(self, directory=None, document=None):
        logger.info("Initializing Corpus object")
        logger.debug(f"Parameter passed: {str([type(i) for i in [directory, document]])}")
        self.index = 0
        self.documents = []
        self.df = pd.DataFrame(columns=['Name', 'Text', 'Label'])
        if directory:
            doc_paths = get_document_paths(directory)
            for path in doc_paths:
                new_doc = Document(path)
                self.add_document(new_doc)
        elif document:
            self.add_document(document)

    def __add__(self, other):
        if type(self) == type(other):
            new_corpus = self
            for document in other.documents:
                new_corpus.add_document(document)
            return new_corpus
        elif type(other) == Document:
            self.add_document(other)
            return self

    def __iter__(self):
        return self

    def __next__(self):
        if self.index < len(self.documents):
            self.index += 1
            return self.documents[self.index - 1]
        else:
            raise StopIteration

    def add_document(self, document):
        logger.info("Adding document to Corpus")
        self.documents.append(document)
        # a_row = pd.Series([document.name, document.text, document.label])
        # row_df = pd.DataFrame([document.name, document.text, document.label])
        self.df = self.df.append({'Name': document.name, 'Text': document.text, 'Label': document.label},
                                 ignore_index=True)
        logger.info("Finished adding document to corpus")

# if __name__ == '__main__':
#     main()
